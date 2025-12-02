import os
from pathlib import Path

def extract_text_from_file(file):
    """
    Extract text from uploaded file.
    Supports: PDF, DOCX, TXT, and images (placeholder for OCR/multimodal AI)
    
    Args:
        file: Django UploadedFile object
    
    Returns:
        str: Extracted text
    """
    file_ext = Path(file.name).suffix.lower()
    
    try:
        if file_ext == '.txt':
            return extract_from_txt(file)
        elif file_ext == '.pdf':
            return extract_from_pdf(file)
        elif file_ext in ['.docx', '.doc']:
            return extract_from_docx(file)
        elif file_ext in ['.jpg', '.jpeg', '.png', '.webp']:
            return extract_from_image(file)
        else:
            return f"Unsupported file type: {file_ext}"
    except Exception as e:
        return f"Error extracting text: {str(e)}"

def extract_from_txt(file):
    """Extract text from TXT file"""
    try:
        file.seek(0)
        content = file.read()
        # Try UTF-8 first, fallback to latin-1
        try:
            return content.decode('utf-8')
        except UnicodeDecodeError:
            return content.decode('latin-1')
    except Exception as e:
        return f"Error reading TXT file: {str(e)}"

def extract_from_pdf(file):
    """Extract text from PDF file using pypdf"""
    try:
        from pypdf import PdfReader
        from io import BytesIO
        
        file.seek(0)
        pdf_file = BytesIO(file.read())
        reader = PdfReader(pdf_file)
        
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n\n"
        
        return text.strip() if text.strip() else "No text found in PDF"
    except ImportError:
        return "PDF support not available. Install 'pypdf' library."
    except Exception as e:
        return f"Error reading PDF: {str(e)}"

def extract_from_docx(file):
    """Extract text from DOCX file using python-docx"""
    try:
        from docx import Document
        from io import BytesIO
        import zipfile
        
        # Reset file pointer to beginning (in case it was read before)
        file.seek(0)
        
        # Read file content once
        try:
            file_content = file.read()
        except Exception as e:
            return f"Error reading file: {str(e)}"
        
        if not file_content:
            return "Error: File is empty"
        
        # Check if it's a valid ZIP file (DOCX is a ZIP archive)
        try:
            docx_file = BytesIO(file_content)
            # Test if it's a valid ZIP
            with zipfile.ZipFile(docx_file, 'r') as zip_ref:
                file_list = zip_ref.namelist()
                # Check for required DOCX files
                if 'word/document.xml' not in file_list:
                    return """Error: This appears to be a ZIP file but not a valid DOCX.

The file is missing 'word/document.xml' which is required for DOCX files.

Please ensure:
- The file was saved as .docx (not .doc)
- The file is not corrupted
- Try re-saving from Microsoft Word"""
        except zipfile.BadZipFile as e:
            return f"""Error: This file is not a valid DOCX file.

Technical details: {str(e)}

Possible causes:
1. The file is corrupted
2. The file is actually a DOC file (older Word format)
3. The file was renamed to .docx but is a different format

Please try:
- Re-saving the file as DOCX in Microsoft Word
- Converting the file to PDF or TXT
- Copying and pasting the text directly"""
        
        # Now extract text - create a fresh BytesIO from the same content
        try:
            docx_file = BytesIO(file_content)
            doc = Document(docx_file)
            
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also try to extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            result = text.strip()
            if not result:
                return "No text found in DOCX. The document might be empty or contain only images."
            return result
        except Exception as e:
            return f"""Error extracting text from DOCX: {str(e)}

This might be due to:
- Unsupported DOCX features
- Incompatible Word version
- Document protection/encryption

Try converting to PDF or TXT format instead."""
    except ImportError as e:
        return f"DOCX support not available. Error: {str(e)}\nInstall 'python-docx' library."
    except Exception as e:
        return f"""Unexpected error reading DOCX: {str(e)}

Please try:
- Converting to PDF or TXT
- Copying and pasting the text directly"""

def extract_from_image(file):
    """
    Placeholder for image text extraction.
    In production, this would use:
    - OCR (pytesseract + Pillow)
    - OR multimodal AI API (e.g., GPT-4 Vision, Gemini Vision)
    """
    return """[Image file uploaded]

NOTE: Image text extraction requires either:
1. OCR setup (pytesseract + Tesseract engine)
2. Multimodal AI API (GPT-4 Vision, Gemini Vision)

For now, please manually type or paste the text from your image below."""
